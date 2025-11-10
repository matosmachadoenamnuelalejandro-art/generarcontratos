from flask import Flask, render_template, request, send_file
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)

# Crear carpeta para documentos si no existe
if not os.path.exists('documentos_generados'):
    os.makedirs('documentos_generados')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generar', methods=['POST'])
def generar_documento():
    try:
        # Obtener datos del formulario
        nombre = request.form['nombre']
        empresa = request.form['empresa']
        servicio = request.form['servicio']
        precio = request.form['precio']
        
        # Crear documento Word
        doc = Document()
        
        # Título
        titulo = doc.add_heading('CONTRATO DE SERVICIOS', 0)
        titulo.alignment = 1  # Centrado
        
        doc.add_paragraph()
        
        # Información de las partes
        doc.add_paragraph(f'ENTRE: {empresa}, en adelante "EL PROVEEDOR"')
        doc.add_paragraph()
        doc.add_paragraph(f'Y: {nombre}, en adelante "EL CLIENTE"')
        
        doc.add_paragraph()
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
        
        # Cláusulas
        doc.add_heading('CLÁUSULAS DEL CONTRATO', level=1)
        
        # Cláusula 1 - Servicio
        doc.add_paragraph()
        doc.add_paragraph(f'1. OBJETO: EL PROVEEDOR se compromete a prestar el servicio de {servicio} a EL CLIENTE.')
        
        # Cláusula 2 - Precio
        doc.add_paragraph()
        doc.add_paragraph(f'2. PRECIO: El valor del servicio es de {precio}.')
        
        # Cláusula 3 - Plazo
        doc.add_paragraph()
        doc.add_paragraph('3. PLAZO: Este contrato tendrá una vigencia de 30 días naturales.')
        
        # Firma y fecha
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y")}')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Firmas
        tabla_firmas = doc.add_table(rows=1, cols=2)
        celdas = tabla_firmas.rows[0].cells
        celdas[0].text = f"____________________\n{empresa}\nPROVEEDOR"
        celdas[1].text = f"____________________\n{nombre}\nCLIENTE"
        
        # Guardar documento
        nombre_archivo = f"contrato_{nombre.replace(' ', '_')}.docx"
        ruta_archivo = os.path.join('documentos_generados', nombre_archivo)
        doc.save(ruta_archivo)
        
        # Datos para mostrar en resultado
        datos = {
            'nombre': nombre,
            'empresa': empresa,
            'servicio': servicio,
            'precio': precio,
            'fecha': datetime.now().strftime("%d/%m/%Y %H:%M")
        }
        
        return render_template('resultado.html', 
                             archivo=nombre_archivo,
                             datos=datos)
    
    except Exception as e:
        return f"Error al generar el documento: {str(e)}"

@app.route('/descargar/<nombre_archivo>')
def descargar(nombre_archivo):
    try:
        return send_file(
            f'documentos_generados/{nombre_archivo}',
            as_attachment=True,
            download_name=nombre_archivo
        )
    except Exception as e:
        return f"Error al descargar: {str(e)}"

if __name__ == '__main__':
    print("✅ Servidor iniciado: http://localhost:5000")
    print("📁 Los documentos se guardan en: documentos_generados/")
    print("🎨 CSS disponible en: http://localhost:5000/static/style.css")
    app.run(debug=True)